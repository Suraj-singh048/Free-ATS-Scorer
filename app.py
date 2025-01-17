import io
import csv
import re
import os
import logging
import json
from typing import Set, Dict, List
from concurrent.futures import ThreadPoolExecutor

import pandas as pd
from flask import Flask, request, render_template, send_file, session, redirect, url_for
from PyPDF2 import PdfReader
from docx import Document

import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

from rapidfuzz import process, fuzz
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

import spacy

from werkzeug.utils import secure_filename
from werkzeug.exceptions import BadRequestKeyError

# Initialize logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s:%(message)s')

# Download essential NLTK data
nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)
nltk.download("stopwords", quiet=True)
nltk.download("wordnet", quiet=True)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
    logging.info("spaCy model loaded successfully.")
except Exception as e:
    logging.error(f"Error loading spaCy model: {e}")
    nlp = None  # Proceed without NER if spaCy model isn't available

app = Flask(__name__)
app.secret_key = "YOUR-SECRET-KEY"  # Required for session usage

# Configuration
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_CONTENT_LENGTH = 5 * 1024 * 1024  # 5 MB
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Configurable thresholds
FUZZY_THRESHOLD = int(os.getenv("FUZZY_THRESHOLD", 80))
FUZZY_THRESHOLD_JD = int(os.getenv("FUZZY_THRESHOLD_JD", 85))

# File paths
SKILL_CSV_PATH = "skill_data/allskills.csv"
SYNONYMS_CSV_PATH = "skill_data/synonyms.csv"
WEIGHTS_CSV_PATH = "skill_data/skill_weights.csv"

# Initialize global variables
ALL_SKILLS: Set[str] = set()
SYNONYMS: Dict[str, List[str]] = {}
SKILL_WEIGHTS: Dict[str, float] = {}
PRECOMPILED_PATTERNS: Dict[str, List[re.Pattern]] = {}

# Initialize NLTK tools
stop_words = set(stopwords.words("english"))
lemmatizer = WordNetLemmatizer()

# Initialize TF-IDF Vectorizer with caching
tfidf_vectorizer = TfidfVectorizer()

# -------------------------------------------------------------------------
# 1. LOAD SKILLS, SYNONYMS & WEIGHTS
# -------------------------------------------------------------------------
def load_csv(filepath: str, required_columns: List[str]) -> pd.DataFrame:
    try:
        df = pd.read_csv(filepath)
        if not set(required_columns).issubset(df.columns):
            raise ValueError(f"CSV at {filepath} is missing required columns: {required_columns}")
        logging.info(f"Loaded CSV: {filepath}")
        return df
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except pd.errors.EmptyDataError:
        logging.error(f"No data: {filepath}")
    except Exception as e:
        logging.error(f"Error loading CSV {filepath}: {e}")
    return pd.DataFrame(columns=required_columns)

def load_synonyms_from_df(df: pd.DataFrame) -> Dict[str, List[str]]:
    synonyms_dict = {}
    for _, row in df.iterrows():
        skill = row["skill"].strip().lower()
        syn_list = []
        if pd.notna(row["synonyms"]) and row["synonyms"].strip():
            syn_list = [s.strip().lower() for s in row["synonyms"].split(";")]
        synonyms_dict[skill] = syn_list
    logging.info("Synonyms loaded successfully.")
    return synonyms_dict

def load_weights_from_df(df: pd.DataFrame) -> Dict[str, float]:
    weights_dict = {}
    for _, row in df.iterrows():
        skill = row["skill"].strip().lower()
        try:
            weight = float(row["weight"])
            weights_dict[skill] = weight
        except ValueError:
            logging.warning(f"Invalid weight for skill '{skill}'; defaulting to 1.0")
            weights_dict[skill] = 1.0
    logging.info("Skill weights loaded successfully.")
    return weights_dict

# Load data
skills_df = load_csv(SKILL_CSV_PATH, ["skill"])
synonyms_df = load_csv(SYNONYMS_CSV_PATH, ["skill", "synonyms"])
weights_df = load_csv(WEIGHTS_CSV_PATH, ["skill", "weight"])

if not skills_df.empty:
    ALL_SKILLS = set(skills_df["skill"].str.lower().unique())
else:
    logging.error("Skills data is empty. Exiting application.")
    exit(1)

if not synonyms_df.empty:
    SYNONYMS = load_synonyms_from_df(synonyms_df)
else:
    logging.warning("Synonyms data is empty.")

if not weights_df.empty:
    SKILL_WEIGHTS = load_weights_from_df(weights_df)
else:
    logging.warning("Skill weights data is empty.")

# Precompile regex patterns for phrase matching
for skill in ALL_SKILLS:
    variants = [skill] + SYNONYMS.get(skill, [])
    PRECOMPILED_PATTERNS[skill] = [re.compile(r"\b" + re.escape(variant) + r"\b", re.IGNORECASE) for variant in variants]

# -------------------------------------------------------------------------
# 2. FILE TEXT EXTRACTION
# -------------------------------------------------------------------------
def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_in_memory(file) -> str:
    """
    Extract text from PDF, DOCX, or TXT file in memory (no saving to disk).
    """
    try:
        filename = file.filename.lower()
        file_data = file.read()

        if filename.endswith(".pdf"):
            text_content = ""
            pdf_reader = PdfReader(io.BytesIO(file_data))
            for page in pdf_reader.pages:
                text_content += (page.extract_text() or "") + " "
            return text_content

        elif filename.endswith(".docx"):
            text_content = ""
            doc = Document(io.BytesIO(file_data))
            for para in doc.paragraphs:
                text_content += para.text + " "
            return text_content

        elif filename.endswith(".txt"):
            return file_data.decode("utf-8", errors="ignore")

        logging.warning(f"Unsupported file type: {filename}")
        return ""

    except Exception as e:
        logging.error(f"Error extracting text from {file.filename}: {e}")
        return ""

# -------------------------------------------------------------------------
# 3. NLTK TOKENIZATION & LEMMATIZATION
# -------------------------------------------------------------------------
def tokenize_and_normalize(text: str) -> str:
    """
    1. Tokenize text.
    2. Keep alpha-only tokens.
    3. Remove stopwords.
    4. Lemmatize.
    5. Return a single string for easier substring and fuzzy matching.
    """
    tokens = word_tokenize(text.lower())
    cleaned = [lemmatizer.lemmatize(t) for t in tokens if t.isalpha() and t not in stop_words]
    return " ".join(cleaned)

# -------------------------------------------------------------------------
# 4. PHRASE & FUZZY MATCHING
# -------------------------------------------------------------------------
def phrase_match(normalized_text: str, skill_pool: Set[str]) -> Set[str]:
    found_skills = set()
    for skill in skill_pool:
        patterns = PRECOMPILED_PATTERNS.get(skill, [])
        for pattern in patterns:
            if pattern.search(normalized_text):
                found_skills.add(skill)
                break
    return found_skills

def fuzzy_match(normalized_text: str, skill_pool: Set[str], threshold: int = 80) -> Set[str]:
    found_skills = set()
    tokens = normalized_text.split()
    all_variants = {variant: skill for skill in skill_pool for variant in [skill] + SYNONYMS.get(skill, [])}

    for token in tokens:
        matches = process.extract(token, all_variants.keys(), scorer=fuzz.ratio, score_cutoff=threshold, limit=1)
        for match, score, _ in matches:
            skill = all_variants[match]
            found_skills.add(skill)
    return found_skills

# -------------------------------------------------------------------------
# 5. TF-IDF COSINE SIMILARITY
# -------------------------------------------------------------------------
def get_cosine_similarity(text1: str, text2: str) -> float:
    try:
        matrix = tfidf_vectorizer.fit_transform([text1, text2])
        sim = cosine_similarity(matrix[0], matrix[1])
        # Flatten the similarity array and convert to float
        return float(sim[0][0])
    except Exception as e:
        logging.error(f"Error computing cosine similarity: {e}")
        return 0.0

# -------------------------------------------------------------------------
# 6. WEIGHTED ATS SCORE
# -------------------------------------------------------------------------
def calculate_weighted_ats(job_skills: Set[str], matched_skills: Set[str]) -> float:
    total_score = sum(SKILL_WEIGHTS.get(skill, 1.0) for skill in job_skills)
    matched_score = sum(SKILL_WEIGHTS.get(skill, 1.0) for skill in matched_skills)
    if total_score == 0:
        return 0.0
    return round((matched_score / total_score) * 100, 2)

# -------------------------------------------------------------------------
# 7. FLASK ROUTES
# -------------------------------------------------------------------------
@app.route("/")
def home():
    return render_template("matchresume.html")  # Original main template

def process_resume(file, jd_relevant_skills: Set[str], job_description: str) -> Dict:
    """
    Process a single resume file (thread-safe).
    Returns the data, but does NOT write to the Flask session.
    """
    filename = secure_filename(file.filename)
    raw_resume_text = extract_text_in_memory(file)
    normalized_resume_text = tokenize_and_normalize(raw_resume_text)

    pm_matches = phrase_match(normalized_resume_text, jd_relevant_skills)
    fm_matches = fuzzy_match(normalized_resume_text, jd_relevant_skills, threshold=FUZZY_THRESHOLD)
    combined_skills = pm_matches.union(fm_matches)

    missing_skills = jd_relevant_skills - combined_skills
    cosim = round(get_cosine_similarity(job_description, raw_resume_text), 3)
    ats_score = calculate_weighted_ats(jd_relevant_skills, combined_skills)

    # Return a dictionary instead of modifying session
    return {
        "filename": filename,
        "raw_resume_text": raw_resume_text,
        "matched_skills": sorted(list(combined_skills)),
        "missing_skills": sorted(list(missing_skills)),
        "ats_score": ats_score,
        "cosine_similarity": cosim
    }

@app.route("/matcher", methods=["POST"])
def matcher():
    try:
        job_description = request.form.get("job_description", "").strip()
        resume_files = request.files.getlist("resumes")

        # [ Validation omitted for brevity ...]

        jd_normalized_text = tokenize_and_normalize(job_description)
        jd_phrase_skills = phrase_match(jd_normalized_text, ALL_SKILLS)
        jd_fuzzy_skills = fuzzy_match(jd_normalized_text, ALL_SKILLS, threshold=FUZZY_THRESHOLD_JD)
        jd_relevant_skills = jd_phrase_skills.union(jd_fuzzy_skills)

        results_data = []
        with ThreadPoolExecutor() as executor:
            futures = [
                executor.submit(process_resume, f, jd_relevant_skills, job_description)
                for f in resume_files
            ]
            for future in futures:
                try:
                    # 1. Get result from the thread
                    resume_data = future.result()

                    # 2. Store any session data you want in the main request context
                    #    (for example, raw text & missing skills)
                    filename = resume_data["filename"]
                    session[filename] = {
                        "raw_text": resume_data["raw_resume_text"],
                        "missing_skills": resume_data["missing_skills"]
                    }

                    # 3. Prepare final "top_resumes" list
                    results_data.append({
                        "filename": filename,
                        "matched_skills": resume_data["matched_skills"],
                        "missing_skills": resume_data["missing_skills"],
                        "ats_score": resume_data["ats_score"],
                        "cosine_similarity": resume_data["cosine_similarity"]
                    })
                except Exception as e:
                    logging.error(f"Error processing a resume: {e}")

        # Sort by ATS score descending
        results_data.sort(key=lambda x: x["ats_score"], reverse=True)

        return render_template(
            "matchresume.html",
            message="Review your ATS results here ==>",
            relevant_skills=sorted(list(jd_relevant_skills)),
            top_resumes=results_data
        )

    except Exception as e:
        logging.error(f"Error in /matcher route: {e}")
        return render_template("matchresume.html", message="An unexpected error occurred. Please try again later.")

@app.route("/download_report", methods=["POST"])
def download_report():
    """
    This route downloads a "summary" report for all resumes in DOCX format.
    (Requires the python-docx library.)
    """
    try:
        results_json = request.form.get("results")
        if not results_json:
            return "No results to download.", 400

        results = json.loads(results_json)

        # Create a new Word document in memory
        doc = Document()
        doc.add_heading('Summary Report', 0)

        # Create a table with headers
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Filename"
        hdr_cells[1].text = "Matched Skills"
        hdr_cells[2].text = "Missing Skills"
        hdr_cells[3].text = "ATS Score"
        hdr_cells[4].text = "Cosine Similarity"

        # Populate table rows
        for result in results:
            row_cells = table.add_row().cells
            row_cells[0].text = result.get("filename", "")
            row_cells[1].text = "; ".join(result.get("matched_skills", []))
            row_cells[2].text = "; ".join(result.get("missing_skills", []))
            row_cells[3].text = str(result.get("ats_score", ""))
            row_cells[4].text = str(result.get("cosine_similarity", ""))

        # Write the doc to a BytesIO object
        temp_stream = io.BytesIO()
        doc.save(temp_stream)
        temp_stream.seek(0)

        # Send the file to the user
        return send_file(
            temp_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='report.docx'
        )

    except Exception as e:
        logging.error(f"Error in /download_report route: {e}")
        return "Failed to generate report.", 500

# -------------------------------------------------------------------------
# 8. NEW: Route to Edit Resume (with missing skills) 
# -------------------------------------------------------------------------
@app.route("/edit_resume/<filename>", methods=["GET"])
def edit_resume(filename):
    """
    Show the user their original resume text and missing skills.
    They can add or modify text here, then click "Download Updated Resume".
    """
    data = session.get(filename)
    if not data:
        return "No data found for that resume.", 404

    raw_text = data["raw_text"]
    missing_skills = data["missing_skills"]

    return render_template(
        "edit_resume.html",
        filename=filename,
        original_resume_text=raw_text,
        missing_skills=missing_skills
    )

# -------------------------------------------------------------------------
# 9. NEW: Route to Download the Updated Resume as .docx
# -------------------------------------------------------------------------
@app.route("/download_updated_resume", methods=["POST"])
def download_updated_resume():
    """
    Receives the edited resume text from the user, then generates a .docx file.
    """
    try:
        from docx import Document

        filename = request.form.get("filename", "")
        updated_text = request.form.get("updated_text", "")

        if not filename:
            return "Filename is missing.", 400

        # Create a docx
        doc = Document()
        doc.add_paragraph(updated_text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Provide the user with a new .docx containing their edits
        updated_filename = f"updated_{filename}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name=updated_filename
        )

    except Exception as e:
        logging.error(f"Error generating updated resume: {e}")
        return "Failed to generate updated resume.", 500

# -------------------------------------------------------------------------
# 10. MAIN
# -------------------------------------------------------------------------
if __name__ == "__main__":
    # Ensure debug is disabled in production
    debug_mode = os.getenv("FLASK_DEBUG", "False").lower() in ['true', '1', 't']
    app.run(debug=debug_mode)
